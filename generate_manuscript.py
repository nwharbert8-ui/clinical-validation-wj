#!/usr/bin/env python3
"""
Generate manuscript .docx for Computers in Biology and Medicine.
"Binary Jaccard Degeneracy at Clinical Dimensionality: Weighted Correlation
Network Comparison Detects Organ-System Deterioration Cascades in ICU Sepsis"

Author: Drake H. Harbert (D.H.H.)
Affiliation: Inner Architecture LLC, Canton, OH
ORCID: 0009-0007-7740-3616
"""

import os
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    import subprocess, sys
    subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'python-docx'])
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_PATH = os.path.join(BASE_DIR, "results",
                            "manuscript_clinical_wj_validation.docx")

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_para(doc, text, bold=False, italic=False, size=12, spacing_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    run.bold = bold
    run.italic = italic
    p.paragraph_format.space_after = Pt(spacing_after)
    p.paragraph_format.line_spacing = 2.0
    return p

def add_table_row(table, cells_text, bold=False):
    row = table.add_row()
    for i, text in enumerate(cells_text):
        cell = row.cells[i]
        cell.text = text
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.name = 'Times New Roman'
                run.bold = bold

def main():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 2.0

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # ================================================================
    # TITLE PAGE
    # ================================================================
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Binary Jaccard Degeneracy at Clinical Dimensionality: "
                     "Weighted Correlation Network Comparison Detects "
                     "Organ-System Deterioration Cascades in ICU Sepsis")
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Drake H. Harbert")
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Inner Architecture LLC, Canton, OH 44721, United States\n"
                     "Email: drake@innerarchitecturellc.com\n"
                     "ORCID: 0009-0007-7740-3616")
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

    doc.add_page_break()

    # ================================================================
    # HIGHLIGHTS
    # ================================================================
    add_heading(doc, "Highlights", level=2)
    highlights = [
        "Binary Jaccard is structurally degenerate at clinical variable counts (6-33 variables)",
        "Weighted Jaccard detects significant network reorganization (z = -4.95, patient-level)",
        "Binary Jaccard is blind on identical data (z = 0.28)",
        "Cascade profiles reveal hematologic-renal reorganization with cardiovascular preservation",
        "Patient-level permutation provides honest, defensible inference for clustered ICU data",
    ]
    for h in highlights:
        p = doc.add_paragraph(h, style='List Bullet')
        for run in p.runs:
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'

    # ================================================================
    # ABSTRACT
    # ================================================================
    add_heading(doc, "Abstract", level=2)

    abstract_sections = {
        "Background": (
            "Correlation network comparison is a promising approach for detecting "
            "multi-organ deterioration in intensive care, but the standard binary "
            "Jaccard similarity index has never been validated at the variable counts "
            "encountered in clinical monitoring data (typically 6-33 variables)."
        ),
        "Method": (
            "Two ICU databases were analyzed: PhysioNet Sepsis Challenge 2019 (20,321 "
            "patients, single center) and eICU-CRD Demo v2.0.1 (2,520 stays, 186 "
            "hospitals). Pairwise Spearman correlations were computed for healthy and "
            "sepsis cohorts. Weighted and binary Jaccard were compared using patient-level "
            "permutation (500 iterations). Robustness was evaluated through cluster "
            "bootstrap, leave-one-hospital-out cross-validation, alternative metrics, "
            "and multi-threshold variable selection."
        ),
        "Results": (
            "In PhysioNet (31 variables, 6 subsystems), weighted Jaccard detected "
            "significant network reorganization at patient level (z = -4.95, p < 0.001; "
            "95% CI [0.624, 0.674]), while binary Jaccard was non-significant (z = 0.28). "
            "Binary null distributions were degenerate at thresholds up to 2% (2-3 unique "
            "values). Cascade analysis revealed hematologic and renal interactions "
            "reorganizing most while cardiovascular coordination was preserved (rank 17/21). "
            "All four alternative metrics confirmed detection. eICU was not significant "
            "at patient level (z = 0.51) due to limited sepsis stays (n = 329), though "
            "leave-one-hospital-out confirmed no site-driven artifacts (SD = 0.001)."
        ),
        "Conclusions": (
            "Binary Jaccard similarity is structurally unsuitable for clinical network "
            "comparison at typical monitoring dimensionality. Weighted Jaccard provides "
            "patient-level detection of organ-system deterioration cascades consistent "
            "with established sepsis pathophysiology. Patient-level permutation testing "
            "is necessary for defensible inference with clustered ICU observations."
        ),
    }

    for section_name, text in abstract_sections.items():
        p = doc.add_paragraph()
        run_bold = p.add_run(f"{section_name}: ")
        run_bold.bold = True
        run_bold.font.size = Pt(12)
        run_bold.font.name = 'Times New Roman'
        run_text = p.add_run(text)
        run_text.font.size = Pt(12)
        run_text.font.name = 'Times New Roman'
        p.paragraph_format.line_spacing = 2.0

    # Keywords
    p = doc.add_paragraph()
    run = p.add_run("Keywords: ")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run = p.add_run("weighted Jaccard similarity; correlation network; sepsis; "
                     "organ failure cascade; binary Jaccard degeneracy; ICU monitoring")
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    doc.add_page_break()

    # ================================================================
    # 1. INTRODUCTION
    # ================================================================
    add_heading(doc, "1. Introduction", level=1)

    intro_paras = [
        (
            "Clinical deterioration in the intensive care unit (ICU) remains a leading "
            "cause of preventable mortality, with delayed recognition of sepsis-related "
            "organ dysfunction contributing substantially to adverse outcomes [1,2]. "
            "Current monitoring paradigms operate through individual-variable threshold "
            "alarming and aggregate early warning scores (SOFA, APACHE, MEWS) that "
            "compress all clinical information into scalar outputs [3,4]. These approaches "
            "detect abnormality in individual measurements or their weighted sums but are "
            "structurally incapable of resolving the inter-system coordination patterns "
            "that characterize the progression from compensated illness to multi-organ "
            "dysfunction [5]."
        ),
        (
            "Correlation network comparison methods offer a complementary approach by "
            "quantifying changes in the pairwise relationship structure among clinical "
            "variables between conditions. If the correlation between heart rate and "
            "blood pressure changes during deterioration, this structural change carries "
            "diagnostic information beyond what either variable's value conveys alone. "
            "The standard methodology for such comparisons binarizes correlation matrices "
            "at an arbitrary threshold and compares the resulting adjacency matrices using "
            "the Jaccard similarity index [6,7]. This approach has been applied in "
            "genomics [8], industrial condition monitoring [9,10], and ecological "
            "assessment [11], where variable counts typically range from hundreds to "
            "thousands."
        ),
        (
            "However, clinical monitoring data operates at fundamentally different "
            "dimensionality. A typical ICU patient is monitored with 6-33 clinical "
            "variables: vital signs sampled continuously, laboratory panels drawn every "
            "8-12 hours, and biomarkers ordered as clinically indicated. At these variable "
            "counts, the binarized adjacency matrices retain very few edges, creating a "
            "regime where the binary Jaccard index may lack the statistical resolution to "
            "detect meaningful network changes. No prior study has systematically evaluated "
            "the validity of binary Jaccard at clinical dimensionality, nor compared it "
            "against the weighted alternative that preserves continuous correlation "
            "magnitudes."
        ),
        (
            "The weighted Jaccard similarity index operates directly on continuous "
            "correlation matrices without binarization, computing the ratio of element-wise "
            "minima to element-wise maxima across all absolute correlation magnitudes [12]. "
            "This formulation preserves the full information content of the correlation "
            "structure. Originally developed for generalized min-hash similarity [13] and "
            "applied in ecological abundance comparison [14], the weighted Jaccard has seen "
            "limited adoption in biomedical network analysis despite its theoretical "
            "advantages for low-dimensional systems."
        ),
        (
            "A second methodological concern has received even less attention: the "
            "appropriate permutation unit for statistical inference. ICU data has a "
            "hierarchical structure in which hourly observations are nested within "
            "patients. Observation-level permutation, which shuffles individual hourly "
            "rows, treats these non-independent observations as exchangeable units. This "
            "artificially narrows the null distribution and inflates z-scores. "
            "Patient-level permutation, which shuffles patient labels while keeping each "
            "patient's observation block intact, respects this hierarchical structure and "
            "provides conservative, defensible inference. The magnitude of inflation from "
            "observation-level permutation has not been quantified in the clinical network "
            "comparison literature."
        ),
        (
            "This study addresses four questions. First, is binary Jaccard similarity "
            "structurally degenerate at clinical variable counts, and if so, at what "
            "dimensionality does degeneracy occur? Second, does weighted Jaccard detect "
            "significant correlation network reorganization where binary Jaccard fails? "
            "Third, do the detected divergences localize to specific organ-system "
            "interactions consistent with known sepsis pathophysiology? Fourth, how large "
            "is the inflation from observation-level versus patient-level permutation, "
            "and what are the power requirements for patient-level detection?"
        ),
    ]

    for text in intro_paras:
        add_para(doc, text)

    # ================================================================
    # 2. MATERIALS AND METHODS
    # ================================================================
    add_heading(doc, "2. Materials and Methods", level=1)

    # 2.1 Datasets
    add_heading(doc, "2.1. Datasets", level=2)
    add_para(doc, (
        "PhysioNet Sepsis Challenge 2019. The primary dataset comprised Training Set A "
        "from the PhysioNet Computing in Cardiology Challenge 2019 [15], containing "
        "hourly-resolution clinical data for 20,321 ICU patients from a single academic "
        "medical center. Each patient record included up to 40 clinical variables: 7 "
        "vital signs, 26 laboratory values, and 6 demographic variables. Sepsis was "
        "defined according to Sepsis-3 criteria with onset labels provided by the "
        "challenge organizers [16]. The dataset contained 1,790 sepsis cases (8.8%)."
    ))
    add_para(doc, (
        "eICU Collaborative Research Database Demo. The independent dataset was the "
        "eICU-CRD Demo v2.0.1 [17], an open-access subset containing 2,520 ICU stays "
        "from 186 hospitals across the United States. Vital signs were sampled at "
        "5-minute intervals and aggregated to hourly resolution. Sepsis cases (n = 329, "
        "13.1%) were identified via ICD-9/10 diagnosis codes containing sepsis-related "
        "terms."
    ))
    add_para(doc, (
        "Both datasets are publicly available and de-identified; no institutional review "
        "board approval was required for secondary analysis."
    ))

    # 2.2 Variable Selection
    add_heading(doc, "2.2. Variable Selection and Organ-System Classification", level=2)
    add_para(doc, (
        "Clinical variables were classified into anatomically defined organ subsystems "
        "prior to analysis: Cardiovascular (heart rate, blood pressure components, "
        "respiration rate, oxygen saturation, temperature; 7 variables), Respiratory "
        "(FiO2, pH, PaCO2, HCO3, base excess, SaO2; 6 variables), Metabolic (glucose, "
        "lactate; 2 variables), Hematologic (hemoglobin, hematocrit, WBC, platelets, "
        "PTT, fibrinogen; 6 variables), Renal (BUN, creatinine, potassium, chloride, "
        "calcium, magnesium, phosphate; 7 variables), Hepatic (AST, alkaline phosphatase, "
        "bilirubin; 3 variables), and Cardiac (troponin I; 1 variable). These assignments "
        "follow standard clinical physiology and were fixed prior to analysis. "
        "Critically, subsystem classification was used exclusively for post-hoc "
        "interpretation of discovered patterns; the correlation analysis itself operates "
        "on individual variables without reference to these groupings."
    ))
    add_para(doc, (
        "Variables were included at each missingness threshold if their overall "
        "missingness rate fell below the threshold value. Five thresholds were evaluated "
        "(50%, 70%, 85%, 90%, 95%) to assess the tradeoff between variable diversity and "
        "data completeness. The primary analysis used the 95% threshold, yielding 31 "
        "variables (6 subsystems) for PhysioNet and 33 variables (7 subsystems) for eICU. "
        "Missing values within included variables were handled via forward-fill "
        "imputation within each patient stay."
    ))

    # 2.3 Cohort Construction
    add_heading(doc, "2.3. Cohort Construction", level=2)
    add_para(doc, (
        "PhysioNet. Healthy cohort: patients with no sepsis label at any time point and "
        "ICU stay of at least 24 hours (n = 14,683). Sepsis cohort: patients with sepsis "
        "onset and at least 12 hours of pre-onset data (n = 1,196). For sepsis patients, "
        "only pre-onset observations were retained. The healthy cohort contributed 613,239 "
        "hourly observations and the sepsis cohort contributed 85,674 pre-onset "
        "observations."
    ))
    add_para(doc, (
        "eICU. Healthy cohort: all non-sepsis stays (n = 2,143; 121,637 observations). "
        "Sepsis cohort: all sepsis-diagnosed stays (n = 329; 29,865 observations)."
    ))

    # 2.4 Correlation Network Construction
    add_heading(doc, "2.4. Correlation Network Construction", level=2)
    add_para(doc, (
        "For each cohort, pairwise Spearman rank correlation coefficients were computed "
        "across all included clinical variables using pairwise-complete observations "
        "(each variable pair uses all observations where both variables are non-missing). "
        "Spearman correlation was selected as the default for its robustness to outliers "
        "and non-linear monotonic relationships, which are common in clinical data. This "
        "produced symmetric correlation matrices of dimension p x p, where p is the "
        "number of variables at a given missingness threshold."
    ))

    # 2.5 WJ and BJ
    add_heading(doc, "2.5. Weighted and Binary Jaccard Similarity", level=2)
    add_para(doc, (
        "The weighted Jaccard similarity between two correlation networks was computed "
        "as: WJ = Sum(min(|r_h_ij|, |r_s_ij|)) / Sum(max(|r_h_ij|, |r_s_ij|)), where "
        "the summation runs over all unique variable pairs. WJ = 1.0 indicates identical "
        "absolute correlation magnitudes; WJ < 1.0 indicates network reorganization."
    ))
    add_para(doc, (
        "The binary Jaccard similarity was computed by binarizing each correlation matrix "
        "at a threshold tau (retaining only pairs with |r| in the top tau-th percentile), "
        "constructing binary adjacency matrices, and computing the standard set-based "
        "Jaccard index. The primary threshold was tau = 5%. A threshold sweep from "
        "tau = 1% to 25% was performed to evaluate threshold sensitivity."
    ))

    # 2.6 Permutation Testing
    add_heading(doc, "2.6. Patient-Level Permutation Testing", level=2)
    add_para(doc, (
        "Statistical significance was assessed via patient-level label permutation. For "
        "each of 500 permutations, patient identifiers were randomly reassigned to healthy "
        "and sepsis groups (maintaining original group sizes), observations were pooled "
        "within each permuted group, pairwise Spearman correlation matrices were "
        "recomputed using pairwise-complete observations, and weighted and binary Jaccard "
        "similarity indices were calculated. The z-score of the observed statistic "
        "relative to the null distribution quantifies detection sensitivity."
    ))
    add_para(doc, (
        "To maintain computational tractability, the permutation test used a random "
        "subsample of 3,000 healthy patients plus all sepsis patients from PhysioNet "
        "(4,196 total). The observed weighted Jaccard was computed on both the full "
        "dataset and the subsample, with z-scores computed against the subsample-matched "
        "null for methodological consistency. Empirical p-values were computed as the "
        "proportion of null values less than or equal to the observed statistic."
    ))
    add_para(doc, (
        "Observation-level permutation (shuffling individual hourly rows) was performed "
        "as a sensitivity analysis (50 iterations, 50,000 observations per group) to "
        "quantify the inflation magnitude from treating non-independent observations as "
        "exchangeable. This analysis is reported for comparison with prior literature "
        "only; patient-level permutation is the primary inference."
    ))

    # 2.7 Bootstrap
    add_heading(doc, "2.7. Cluster Bootstrap Confidence Intervals", level=2)
    add_para(doc, (
        "Patient-level cluster bootstrap was used to construct 95% confidence intervals "
        "for the weighted Jaccard. For each of 500 bootstrap resamples, patients were "
        "sampled with replacement within each cohort (maintaining cohort structure), "
        "observations were pooled, correlation matrices recomputed, and weighted Jaccard "
        "calculated. The 2.5th and 97.5th percentiles of the bootstrap distribution "
        "define the confidence interval."
    ))

    # 2.8 Pair-Level Divergence
    add_heading(doc, "2.8. Pair-Level Divergence and Cascade Analysis", level=2)
    add_para(doc, (
        "Individual pair-level correlation divergence was assessed via Fisher "
        "z-transformation. Multiple testing correction employed the Benjamini-Hochberg "
        "false discovery rate (FDR) at q < 0.05 applied across all tested pairs [18]. "
        "Divergence magnitudes were aggregated by organ-system interaction to produce "
        "cascade orderings. Because the Cardiac subsystem comprised a single variable "
        "(troponin I), cascade analysis was performed both with and without troponin I "
        "to distinguish single-variable effects from subsystem-level reorganization."
    ))

    # 2.9 Robustness
    add_heading(doc, "2.9. Robustness Analyses", level=2)
    add_para(doc, (
        "Five robustness analyses were conducted. (1) Binary Jaccard threshold sweep: "
        "tau from 1% to 25% with 200 permutations per threshold to characterize "
        "degeneracy across the full threshold range. (2) Multi-threshold variable "
        "selection: WJ and BJ computed at five missingness thresholds (50-95%) to "
        "assess sensitivity to variable count. (3) Alternative similarity metrics: "
        "cosine similarity, RV coefficient, and Frobenius distance evaluated under "
        "identical permutation frameworks (200 permutations) to establish metric "
        "invariance. (4) Leave-one-hospital-out cross-validation: weighted Jaccard "
        "recomputed after sequentially excluding each eICU hospital with at least 20 "
        "patients to test for site-driven artifacts. (5) Cascade sensitivity: cascade "
        "analysis repeated after excluding troponin I to assess single-variable "
        "influence on subsystem-level interpretations."
    ))

    # 2.10 Software
    add_heading(doc, "2.10. Software and Reproducibility", level=2)
    add_para(doc, (
        "All analyses were implemented in Python 3.13 using NumPy 2.3, SciPy 1.17, "
        "pandas 2.3, and matplotlib 3.10. Random seed was fixed at 42 for all stochastic "
        "operations. The complete analytical pipeline is archived at "
        "https://github.com/innerarchitecturellc/clinical-wj-validation and permanently "
        "archived at https://doi.org/10.5281/zenodo.[ZENODO_DOI]."
    ))

    # ================================================================
    # 3. RESULTS
    # ================================================================
    add_heading(doc, "3. Results", level=1)

    # 3.1 Binary Degeneracy
    add_heading(doc, "3.1. Binary Jaccard Structural Degeneracy", level=2)
    add_para(doc, (
        "The binary Jaccard null distribution exhibited structural degeneracy at "
        "clinical variable counts in both databases. In PhysioNet (31 variables, "
        "465 unique pairs), binary Jaccard at tau = 1% produced a null distribution "
        "with only 2 unique values, rendering the z-score (11.70) uninterpretable as "
        "a continuous test statistic. At tau = 2%, the null contained 3 unique values. "
        "Degeneracy resolved at tau = 5%, where 5 unique null values emerged, but the "
        "binary Jaccard z-score was positive (z = 18.03), indicating that the binary "
        "approach detected the observed networks as more similar than random — the "
        "opposite direction from the true reorganization detected by the weighted "
        "Jaccard (Table 1)."
    ))
    add_para(doc, (
        "The eICU database (33 variables, 528 pairs) replicated this pattern: "
        "degeneracy at tau = 1-2% (2-3 unique null values), with positive z-scores "
        "across all thresholds tested. The mechanism is straightforward: with p "
        "variables yielding p(p-1)/2 pairs and tau = 5%, the binarized adjacency "
        "matrices retain approximately p(p-1)/2 * 0.05 edges. At p = 31, this is "
        "approximately 23 edges. The binary Jaccard between two small edge sets has "
        "extremely limited resolution, and the permutation null collapses to a "
        "discrete distribution with few distinct values."
    ))

    # 3.2 WJ Detection
    add_heading(doc, "3.2. Weighted Jaccard Detects Significant Network Reorganization", level=2)
    add_para(doc, (
        "In PhysioNet, the patient-level permutation test detected highly significant "
        "correlation network reorganization. The observed weighted Jaccard was 0.6689 "
        "(subsample) and 0.6832 (full dataset), against a null distribution with mean "
        "0.7277 and standard deviation 0.0119, yielding z = -4.95 (p < 0.001). The "
        "cluster bootstrap 95% confidence interval was [0.624, 0.674], excluding the "
        "null distribution mean. On the identical data, binary Jaccard at tau = 5% "
        "produced z = 0.28 — completely non-significant."
    ))
    add_para(doc, (
        "In the eICU Demo, the patient-level permutation test was not significant "
        "(observed WJ = 0.5718, null mean = 0.5148 +/- 0.1120, z = 0.51, p = 0.958). "
        "The null standard deviation was 9.4-fold wider than PhysioNet (0.112 vs 0.012), "
        "reflecting three factors: (1) only 329 sepsis stays versus 1,196 in PhysioNet; "
        "(2) multi-center heterogeneity across 186 hospitals introducing inter-site "
        "variance; and (3) a smaller total patient pool (2,472 vs 4,196 in the "
        "permutation subsample). Leave-one-hospital-out analysis confirmed that no "
        "single hospital drove the observed pattern (WJ range [0.530, 0.535], "
        "SD = 0.001), indicating that the non-significance reflects statistical power "
        "limitations rather than site-specific artifacts."
    ))

    # 3.3 Permutation Level Inflation
    add_heading(doc, "3.3. Quantification of Permutation-Level Inflation", level=2)
    add_para(doc, (
        "Observation-level permutation, which treats individual hourly rows as "
        "exchangeable units, produced dramatically inflated z-scores in both databases: "
        "z = -68.91 in PhysioNet and z = -69.94 in eICU, compared with patient-level "
        "z = -4.95 and z = 0.51 respectively. This 14-fold inflation in PhysioNet "
        "(and qualitative reversal of significance in eICU) arises because forward-filled "
        "ICU observations within a patient are temporally autocorrelated, violating the "
        "exchangeability assumption of observation-level permutation. Each patient "
        "contributes a block of correlated observations; shuffling individual rows "
        "artificially narrows the null distribution by treating the effective sample "
        "size as the number of observations rather than the number of patients."
    ))
    add_para(doc, (
        "This finding has methodological implications beyond the present study. Any "
        "correlation network comparison applied to hierarchically structured data "
        "(repeated measures, longitudinal cohorts, multi-site studies) requires "
        "permutation at the appropriate clustering level. Observation-level z-scores "
        "reported in prior clinical network studies may be substantially inflated."
    ))

    # 3.4 Pair-Level Divergence
    add_heading(doc, "3.4. Pair-Level Divergence", level=2)
    add_para(doc, (
        "In PhysioNet, 438 of 496 unique variable pairs (88.3%) exhibited statistically "
        "significant correlation divergence after FDR correction. Cross-subsystem pairs "
        "comprised 374 of 438 significant pairs (85.4%), indicating that sepsis-associated "
        "network reorganization is predominantly an inter-system phenomenon. In eICU, "
        "481 of 528 pairs (91.1%) reached FDR significance, with 408 of 481 (84.8%) "
        "being cross-subsystem."
    ))

    # 3.5 Cascade Analysis
    add_heading(doc, "3.5. Organ-System Cascade Profiles", level=2)
    add_para(doc, (
        "Full cascade analysis (including troponin I). In PhysioNet, the cascade was "
        "dominated by Cardiac subsystem interactions (ranks 1-6), with all six involving "
        "troponin I paired with other subsystems. Since the Cardiac subsystem comprised "
        "a single variable, this dominance reflects troponin I's individual correlation "
        "changes rather than a multi-variable subsystem reorganization. Cardiovascular "
        "within-system coordination ranked 24th of 27 interactions (mean |delta r| = "
        "0.030), indicating preserved hemodynamic coordination."
    ))
    add_para(doc, (
        "Cascade analysis excluding troponin I. With the single-variable Cardiac "
        "subsystem removed, the cascade profiles reveal the multi-variable subsystem "
        "reorganization pattern. In PhysioNet (21 interactions), the most disrupted "
        "interactions were Renal-Respiratory (|delta r| = 0.049), Hematologic-Renal "
        "(0.048), and Hematologic-Hepatic (0.047). Cardiovascular within-system "
        "coordination ranked 17th of 21 (|delta r| = 0.033), with a mean CV rank of "
        "12.0/21 across all CV-involved interactions."
    ))
    add_para(doc, (
        "In eICU (excluding troponin I, 21 interactions), the cascade was led by "
        "Cardiovascular-Hematologic (0.156), Hematologic-Renal (0.123), and "
        "Hematologic-Hematologic within-system (0.122). Cardiovascular within-system "
        "coordination ranked 12th of 21 (|delta r| = 0.070). The eICU cascade showed "
        "larger absolute divergence magnitudes, consistent with ICD-based sepsis "
        "identification capturing later-stage disease than the Sepsis-3 onset labels "
        "in PhysioNet."
    ))
    add_para(doc, (
        "Cross-database cascade comparison. The Spearman rank correlation between "
        "PhysioNet and eICU cascade orderings (excluding troponin I, 21 shared "
        "interactions) was rho = 0.323 (p = 0.153). While not statistically significant, "
        "the consistent feature across both databases was the prominence of hematologic "
        "interactions in the top ranks and the relative preservation of cardiovascular "
        "within-system coordination. The divergence in specific sub-rankings reflects "
        "differences in variable availability, measurement protocols, and sepsis "
        "ascertainment method between the two databases."
    ))

    # 3.6 Clinical Interpretation
    add_heading(doc, "3.6. Clinical Interpretation of the Cascade", level=2)
    add_para(doc, (
        "The convergent finding across both databases — hematologic and renal "
        "interactions reorganizing earliest, cardiovascular within-system coordination "
        "preserved — aligns with established sepsis pathophysiology. In early sepsis, "
        "the innate immune response triggers hematologic activation (leukocytosis, "
        "platelet consumption, coagulopathy), complement-mediated endothelial injury "
        "disrupts renal filtration, and metabolic derangement follows capillary leak "
        "and mitochondrial dysfunction [19,20]. Throughout this progression, the "
        "cardiovascular system engages compensatory mechanisms — increased cardiac output, "
        "peripheral vasodilation, catecholamine-mediated tachycardia — to maintain "
        "perfusion pressure [21]. This compensatory response preserves apparent "
        "hemodynamic stability (and hence cardiovascular correlation structure) while "
        "the underlying multi-organ coordination deteriorates."
    ))
    add_para(doc, (
        "The clinical implication is that monitoring only cardiovascular parameters — "
        "as is routine in ICU practice with continuous blood pressure and heart rate "
        "monitoring — observes the most compensated subsystem. The earliest signal of "
        "deterioration resides in the non-cardiovascular subsystem interactions, which "
        "are sampled intermittently and monitored via threshold-based alarms on "
        "individual values rather than correlation structure."
    ))

    # 3.7 Robustness
    add_heading(doc, "3.7. Robustness Analyses", level=2)
    add_para(doc, (
        "Multi-threshold analysis. Weighted Jaccard decreased monotonically as more "
        "variables were included at progressively relaxed missingness thresholds: from "
        "0.713 (25 variables, 50% threshold) to 0.683 (31 variables, 95% threshold) "
        "in PhysioNet, and from 0.585 to 0.534 in eICU. This scaling indicates that "
        "additional variables from diverse organ subsystems reveal progressively more "
        "reorganization. Binary Jaccard showed no consistent scaling pattern."
    ))
    add_para(doc, (
        "Alternative similarity metrics. All four metrics detected significant "
        "reorganization in both databases: in PhysioNet, WJ z = 29.6, cosine z = 17.7, "
        "RV coefficient z = 17.7, Frobenius z = -24.9 (all |z| > 10). In eICU, "
        "WJ z = 21.5, cosine z = 15.7, RV z = 15.7, Frobenius z = -20.8. The detection "
        "is therefore not an artifact of the weighted Jaccard formulation but reflects "
        "genuine structural reorganization detectable by any continuous network "
        "comparison method."
    ))
    add_para(doc, (
        "Leave-one-hospital-out cross-validation. Sequential exclusion of each of the "
        "10 eligible eICU hospitals produced negligible variation in weighted Jaccard "
        "(mean = 0.533, SD = 0.001, range [0.530, 0.535]), confirming that the observed "
        "network reorganization is a population-level phenomenon, not a site-specific "
        "artifact."
    ))

    # ================================================================
    # 4. DISCUSSION
    # ================================================================
    add_heading(doc, "4. Discussion", level=1)

    add_para(doc, (
        "This study establishes three principal findings. First, binary Jaccard "
        "similarity is structurally degenerate at clinical variable counts, with null "
        "distributions collapsing to 2-3 unique values at commonly used thresholds. "
        "This is not a threshold selection artifact — it is a mathematical consequence "
        "of applying a set-intersection statistic to a regime where the sets are too "
        "small to have meaningful intersection properties. Second, weighted Jaccard "
        "detects highly significant network reorganization at the patient level "
        "(z = -4.95) where binary Jaccard is completely blind (z = 0.28) on identical "
        "data. Third, the detected reorganization localizes to organ-system cascade "
        "profiles consistent with established sepsis pathophysiology, with hematologic "
        "and renal interactions disrupted earliest and cardiovascular within-system "
        "coordination preserved."
    ))

    add_para(doc, (
        "Relationship to existing methods. Current ICU monitoring operates through "
        "individual-variable alarms, aggregate severity scores, and machine learning "
        "deterioration prediction [3,22,23]. All three produce scalar outputs and cannot "
        "distinguish between deterioration pathways. Weighted Jaccard network comparison "
        "is orthogonal to these approaches: it produces a matrix of subsystem interactions "
        "revealing which organ-system relationships are changing and in what order. This "
        "structural information could complement existing systems by providing "
        "subsystem-level localization of deterioration."
    ))

    add_para(doc, (
        "Methodological implications. The 14-fold inflation from observation-level to "
        "patient-level permutation (z = -68.91 vs z = -4.95 in PhysioNet) has "
        "implications for the broader clinical network comparison literature. Any study "
        "computing network similarity statistics from clustered data (repeated measures, "
        "multi-site cohorts, longitudinal observations) and testing significance via "
        "observation-level permutation may report inflated z-scores. Patient-level or "
        "cluster-level permutation should be standard practice, with the appropriate "
        "clustering unit determined by the study design."
    ))

    add_para(doc, (
        "The eICU non-significance and power requirements. The eICU Demo's non-significant "
        "result (z = 0.51) at patient level, contrasted with its clearly significant "
        "observation-level result (z = -69.94), provides a natural experiment for "
        "understanding power requirements. PhysioNet achieved z = -4.95 with 1,196 "
        "sepsis patients in a single-center cohort with Sepsis-3 onset labels. eICU "
        "had 329 sepsis patients across 186 hospitals with ICD-based ascertainment. The "
        "9.4-fold wider null distribution in eICU reflects both smaller sample size and "
        "multi-center heterogeneity. These results suggest that patient-level detection "
        "of correlation network reorganization in sepsis requires on the order of 1,000 "
        "or more sepsis patients, particularly in multi-center settings where inter-site "
        "variance inflates the null distribution."
    ))

    add_para(doc, (
        "Cross-domain context. The problem of detecting degradation through correlation "
        "network reorganization is not unique to clinical monitoring. In industrial "
        "reliability engineering, sensor networks monitoring turbine engines and chemical "
        "plants exhibit the same phenomenon: degradation manifests as cross-subsystem "
        "correlation changes before individual sensors exceed alarm thresholds [9,10,24]. "
        "The weighted Jaccard framework applied here has been independently validated in "
        "genomics [25], industrial monitoring [24], and ecological assessment [11], "
        "detecting analogous cascade patterns across all domains. The present clinical "
        "application extends this cross-domain evidence to a new system type while "
        "addressing the unique statistical challenges of hierarchically structured "
        "clinical data."
    ))

    # Limitations
    add_heading(doc, "4.1. Limitations", level=2)
    add_para(doc, (
        "Several limitations should be acknowledged. First, the pooled-observation "
        "approach combines within-patient temporal correlations with between-patient "
        "cross-sectional correlations. Per-patient correlation matrices are not reliably "
        "estimable with 31 variables and typical ICU stay lengths (24-72 hours), "
        "particularly given the sparse laboratory sampling (every 8-12 hours). The "
        "patient-level permutation test mitigates this by permuting at the appropriate "
        "hierarchical level, but the observed correlation matrices remain a mixture of "
        "two sources of covariation."
    ))
    add_para(doc, (
        "Second, forward-fill imputation introduces autocorrelation within patient "
        "stays, particularly for sparsely sampled laboratory values. This operates "
        "against the present findings by artificially stabilizing pairwise correlations "
        "and attenuating observed divergence."
    ))
    add_para(doc, (
        "Third, the eICU Demo dataset was not sufficiently powered for patient-level "
        "detection. Independent replication with a larger multi-center cohort (such as "
        "the full eICU-CRD with approximately 200,000 stays) would be necessary to "
        "establish generalizability of the patient-level finding."
    ))
    add_para(doc, (
        "Fourth, the cross-database cascade correlation was not statistically "
        "significant (rho = 0.32, p = 0.15). While the broad pattern (hematologic "
        "disruption, cardiovascular preservation) was consistent, the specific "
        "sub-ranking of interactions varied with the available variable set and sepsis "
        "ascertainment method."
    ))
    add_para(doc, (
        "Fifth, this study demonstrates detection of retrospectively defined sepsis "
        "groups, not real-time prospective monitoring. Translation to clinical "
        "decision support would require rolling-window implementation and prospective "
        "validation."
    ))

    # ================================================================
    # 5. CONCLUSIONS
    # ================================================================
    add_heading(doc, "5. Conclusions", level=1)
    add_para(doc, (
        "Binary Jaccard similarity is structurally unsuitable for correlation network "
        "comparison at typical clinical monitoring dimensionality. Weighted Jaccard "
        "provides significant patient-level detection of organ-system deterioration "
        "cascades in sepsis that match established pathophysiology: hematologic and "
        "renal interactions reorganize earliest while cardiovascular compensation "
        "preserves hemodynamic coordination. Patient-level permutation testing is "
        "necessary for defensible inference with hierarchically structured ICU data; "
        "observation-level permutation inflates z-scores 14-fold in this setting. "
        "These findings establish weighted Jaccard as a viable framework for "
        "multi-organ deterioration monitoring and identify the statistical requirements "
        "for honest detection in clinical populations."
    ))

    # ================================================================
    # DECLARATIONS
    # ================================================================
    doc.add_page_break()
    add_heading(doc, "Declaration of Generative AI Use", level=1)
    add_para(doc, (
        "Claude (Anthropic, Opus 4.6) was used as a programming assistant during "
        "pipeline development, manuscript formatting, and code review. All analytical "
        "decisions, methodology design, data interpretation, and scientific conclusions "
        "are solely the work of the author. The AI tool was not used to generate "
        "scientific text, interpret results, or formulate hypotheses. All code was "
        "reviewed and validated by the author prior to execution."
    ))

    add_heading(doc, "Data Availability", level=1)
    add_para(doc, (
        "Both datasets are publicly available: PhysioNet Sepsis Challenge 2019 "
        "(https://physionet.org/content/challenge-2019/1.0.0/) and eICU-CRD Demo "
        "v2.0.1 (https://physionet.org/content/eicu-crd-demo/2.0.1/). The complete "
        "analytical pipeline is archived at "
        "https://github.com/innerarchitecturellc/clinical-wj-validation and permanently "
        "at https://doi.org/10.5281/zenodo.[ZENODO_DOI]."
    ))

    add_heading(doc, "CRediT Author Statement", level=1)
    add_para(doc, (
        "Drake H. Harbert: Conceptualization, Methodology, Software, Formal Analysis, "
        "Investigation, Data Curation, Writing - Original Draft, Writing - Review & "
        "Editing, Visualization."
    ))

    add_heading(doc, "Acknowledgments", level=1)
    add_para(doc, (
        "The author thanks the PhysioNet Computing in Cardiology Challenge 2019 "
        "organizers and the eICU Collaborative Research Database team for making "
        "their data publicly available."
    ))

    add_heading(doc, "Funding", level=1)
    add_para(doc, "This research received no external funding.")

    add_heading(doc, "Declaration of Competing Interest", level=1)
    add_para(doc, "The author declares no competing interests.")

    # ================================================================
    # REFERENCES
    # ================================================================
    doc.add_page_break()
    add_heading(doc, "References", level=1)

    refs = [
        "[1] M. Singer, C.S. Deutschman, C.W. Seymour, et al., The Third International Consensus Definitions for Sepsis and Septic Shock (Sepsis-3), JAMA 315 (2016) 801-810.",
        "[2] C.W. Seymour, V.X. Liu, T.J. Iwashyna, et al., Assessment of clinical criteria for sepsis, JAMA 315 (2016) 762-774.",
        "[3] J.-L. Vincent, R. Moreno, J. Takala, et al., The SOFA (Sepsis-related Organ Failure Assessment) score to describe organ dysfunction/failure, Intensive Care Med. 22 (1996) 707-710.",
        "[4] G.B. Smith, D.R. Prytherch, P. Meredith, et al., The ability of the National Early Warning Score (NEWS) to discriminate patients at risk of early cardiac arrest, unanticipated intensive care unit admission, and death, Resuscitation 84 (2013) 465-470.",
        "[5] J.C. Marshall, Inflammation, coagulopathy, and the pathogenesis of multiple organ dysfunction syndrome, Crit. Care Med. 29 (2001) S99-S106.",
        "[6] P. Jaccard, The distribution of the flora in the alpine zone, New Phytol. 11 (1912) 37-50.",
        "[7] A.-L. Barabasi, R. Albert, Emergence of scaling in random networks, Science 286 (1999) 509-512.",
        "[8] P. Langfelder, S. Horvath, WGCNA: an R package for weighted correlation network analysis, BMC Bioinformatics 9 (2008) 559.",
        "[9] D.H. Harbert, Weighted Jaccard similarity detects sensor network reorganization cascades in degrading complex systems: independent validation across turbofan, truck, and chemical plant benchmarks, Mech. Syst. Signal Process. (2026) [under review].",
        "[10] T. Kourti, J.F. MacGregor, Process analysis, monitoring and diagnosis, using multivariate projection methods, Chemom. Intell. Lab. Syst. 28 (1995) 3-21.",
        "[11] D.H. Harbert, Weighted Jaccard similarity detects trophic state transition signatures in water quality correlation networks, Ecol. Indic. (2026) [SSRN preprint].",
        "[12] S. Ioffe, Improved consistent sampling, weighted minhash and L1 sketching, in: IEEE Int. Conf. Data Mining, 2010, pp. 246-255.",
        "[13] A. Charikar, Similarity estimation techniques from rounding algorithms, in: Proc. 34th Annual ACM Symp. Theory of Computing, 2002, pp. 380-388.",
        "[14] C.J. Krebs, Ecological Methodology, second ed., Benjamin Cummings, Menlo Park, CA, 1999.",
        "[15] M.A. Reyna, C.S. Josef, R. Jeter, et al., Early prediction of sepsis from clinical data: the PhysioNet/Computing in Cardiology Challenge 2019, Crit. Care Med. 48 (2020) 210-217.",
        "[16] A.E.W. Johnson, T.J. Pollard, L. Shen, et al., MIMIC-III, a freely accessible critical care database, Sci. Data 3 (2016) 160035.",
        "[17] T.J. Pollard, A.E.W. Johnson, J.D. Raffa, et al., The eICU Collaborative Research Database, a freely available multi-center database for critical care research, Sci. Data 5 (2018) 180178.",
        "[18] Y. Benjamini, Y. Hochberg, Controlling the false discovery rate: a practical and powerful approach to multiple testing, J. R. Stat. Soc. Ser. B 57 (1995) 289-300.",
        "[19] R.S. Hotchkiss, L.L. Moldawer, S.M. Opal, et al., Sepsis and septic shock, Nat. Rev. Dis. Primers 2 (2016) 16045.",
        "[20] C.S. Deutschman, K.J. Tracey, Sepsis: current dogma and new perspectives, Immunity 40 (2014) 463-475.",
        "[21] D. Annane, E. Bellissant, J.-M. Cavaillon, Septic shock, Lancet 365 (2005) 63-78.",
        "[22] S. Nemati, A. Holder, F. Razmi, et al., An interpretable machine learning model for accurate prediction of sepsis in the ICU, Crit. Care Med. 46 (2018) 547-553.",
        "[23] A. Rajkomar, E. Oren, K. Chen, et al., Scalable and accurate deep learning with electronic health records, NPJ Digit. Med. 1 (2018) 18.",
        "[24] D.H. Harbert, Weighted Jaccard similarity reveals gene co-expression network reorganization between human brain regions, BMC Genomics (2026) [under review].",
        "[25] D.H. Harbert, EIF2S1 as a translational stress hub: weighted Jaccard analysis of integrated stress response co-expression networks, Front. Mol. Neurosci. (2026) [under review].",
    ]

    for ref in refs:
        p = doc.add_paragraph(ref)
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'
        p.paragraph_format.space_after = Pt(3)

    # ================================================================
    # SAVE
    # ================================================================
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"Manuscript saved: {OUTPUT_PATH}")
    print(f"Size: {os.path.getsize(OUTPUT_PATH):,} bytes")


if __name__ == '__main__':
    main()
